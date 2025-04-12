import wikipedia
from docx import Document
import datetime
import os

def do_research(topic, page_limit=10):
    """Generate a detailed research document on a topic"""
    try:
        wikipedia.set_lang("en")
        summary = wikipedia.summary(topic, sentences=2)

        pages = wikipedia.page(topic)
        full_content = pages.content

        # Split into smaller chunks based on newlines
        chunks = full_content.split('\n\n')

        # Limit content length based on page count estimate
        approx_chars_per_page = 1800
        max_chars = page_limit * approx_chars_per_page
        content = "\n\n".join(chunks)
        content = content[:max_chars]

        # Create a Word doc
        doc = Document()
        doc.add_heading(f"Research on {topic.title()}", 0)
        doc.add_paragraph(f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph("\nSummary:\n" + summary)
        doc.add_paragraph("\nDetailed Content:\n")
        doc.add_paragraph(content)

        filename = f"{topic.lower().replace(' ', '_')}_research.docx"
        filepath = os.path.join(os.path.expanduser("~"), "Desktop", filename)
        doc.save(filepath)

        return f"Research saved successfully to Desktop as {filename}."

    except Exception as e:
        return f"Error during research: {str(e)}"
